from docx import Document
from datetime import date
from django.db.models import Sum
from principal.utils import perfil
from django.http import HttpResponse
from django.shortcuts import render, redirect
from principal.models import Receita, Despesa


# RETORNA DADOS DO BALANÇO PARA RELATÓRIO COMPLETO
def balanco(request):
    usuario_id = request.session.get('usuario_id')

    if request.method == "GET":
        if usuario_id:
            contexto = perfil(request)

            receitas = Receita.objects.filter(id_usuario=usuario_id, status='efetuada').values('descricao', 'valor',
                                                                                            'data_receita')
            despesas = Despesa.objects.filter(id_usuario=usuario_id, status='efetuada').values('descricao', 'valor',
                                                                                            'data_despesa')

            total_receitas = receitas.aggregate(Sum('valor'))['valor__sum'] or 0
            total_despesas = despesas.aggregate(Sum('valor'))['valor__sum'] or 0

            total_ativos = total_receitas
            total_passivos = total_despesas
            patrimonio_liquido = total_ativos - total_passivos

            contexto.update({
                'receitas': receitas,
                'despesas': despesas,
                'total_receitas': total_receitas,
                'total_despesas': total_despesas,
                'total_ativos': total_ativos,
                'total_passivos': total_passivos,
                'patrimonio_liquido': patrimonio_liquido,
            })

            if 'download' in request.GET:
                return gerar_relatorio_word_balanco(receitas, despesas, total_receitas, total_despesas, total_ativos,
                                            total_passivos, patrimonio_liquido)

            return render(request, 'relatorio/relatorio1.html', contexto)
        else:
            return redirect('login')


# RELATÓRIO WORD BALNAÇO
def gerar_relatorio_word_balanco(receitas, despesas, total_receitas, total_despesas, total_ativos, total_passivos, patrimonio_liquido):
    document = Document()
    document.add_heading('Relatório de Balanço Patrimonial', 0)

    document.add_heading('Ativos', level=1)
    for receita in receitas:
        document.add_paragraph(f"{receita['descricao']} - ${receita['valor']} em {receita['data_receita']}")

    document.add_heading('Passivos', level=1)
    for despesa in despesas:
        document.add_paragraph(f"{despesa['descricao']} - ${despesa['valor']} em {despesa['data_despesa']}")

    document.add_heading('Resumo', level=1)
    document.add_paragraph(f"Total de Ativos: ${total_receitas}")
    document.add_paragraph(f"Total de Passivos: ${total_despesas}")
    document.add_paragraph(f"Total de Ativos: ${total_ativos}")
    document.add_paragraph(f"Total de Passivos: ${total_passivos}")
    document.add_paragraph(f"Patrimônio Líquido: ${patrimonio_liquido}")

    response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    response['Content-Disposition'] = 'attachment; filename=Relatorio_de_Balanco_Patrimonial.docx'
    document.save(response)
    return response



# RETORNA DADOS DE CONTAS PARA RELATÓRIO COMPLETO
def contas(request):
    usuario_id = request.session.get('usuario_id')

    if request.method == "GET":
        if usuario_id:
            contexto = perfil(request)

            despesas = Despesa.objects.filter(id_usuario=usuario_id, status='nao_efetuada').values('descricao', 'valor',
                                                                                             'data_despesa')
            receitas_a_receber = Receita.objects.filter(id_usuario=usuario_id, status='nao_efetuada').values('descricao',
                                                                                                       'valor',
                                                                                                       'data_receita')

            despesas_atrasadas = [despesa for despesa in despesas if despesa['data_despesa'] < date.today()]

            total_contas_a_pagar = sum(despesa['valor'] for despesa in despesas)
            total_receitas_a_receber = sum(receita['valor'] for receita in receitas_a_receber)
            total_despesas_atrasadas = sum(despesa['valor'] for despesa in despesas_atrasadas)

            contexto.update({
                'despesas_a_pagar': despesas,
                'receitas_a_receber': receitas_a_receber,
                'total_contas_a_pagar': total_contas_a_pagar,
                'total_receitas_a_receber': total_receitas_a_receber,
                'total_despesas_atrasadas': total_despesas_atrasadas,
            })

            if 'download' in request.GET:
                return gerar_relatorio_word_contas(despesas_atrasadas, receitas_a_receber, total_contas_a_pagar,
                                            total_receitas_a_receber, total_despesas_atrasadas)

            return render(request, 'relatorio/relatorio2.html', contexto)
        else:
            return redirect('login')


# RELATÓRIO WORD CONTAS
def gerar_relatorio_word_contas(despesas_atrasadas, receitas_a_receber, total_contas_a_pagar, total_receitas_a_receber, total_despesas_atrasadas):
    document = Document()
    document.add_heading('Relatório de Contas', 0)

    document.add_heading('Despesas Atrasadas', level=1)
    for despesa in despesas_atrasadas:
        document.add_paragraph(f"{despesa['descricao']} - ${despesa['valor']} data da despesa {despesa['data_despesa']}")

    document.add_heading('Receitas a Receber', level=1)
    for receita in receitas_a_receber:
        document.add_paragraph(f"{receita['descricao']} - ${receita['valor']} data esperada {receita['data_receita']}")

    document.add_heading('Resumo', level=1)
    document.add_paragraph(f"Total de Contas a Pagar (Despesas): ${total_contas_a_pagar}")
    document.add_paragraph(f"Total de Receitas a Receber: ${total_receitas_a_receber}")
    document.add_paragraph(f"Total de Despesas Atrasadas: ${total_despesas_atrasadas}")

    response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    response['Content-Disposition'] = 'attachment; filename=Relatorio_de_Contas.docx'
    document.save(response)
    return response


# RETORNA DADOS DE GASTOS PARA RELATÓRIO COMPLETO
def gastos(request):
    usuario_id = request.session.get('usuario_id')

    if request.method == "GET":
        if usuario_id:
            contexto = perfil(request)

            despesas = Despesa.objects.filter(id_usuario=usuario_id, status='efetuada').values('descricao', 'valor',
                                                                                            'data_despesa')

            total_gastos = despesas.aggregate(Sum('valor'))['valor__sum'] or 0
            maior_gasto = max(despesas, key=lambda x: x['valor']) if despesas else None
            maior_mes_gasto = max(despesas, key=lambda x: x['data_despesa'].strftime('%Y-%m')) if despesas else None

            contexto.update({
                'despesas': despesas,
                'total_gastos': total_gastos,
                'maior_gasto': maior_gasto,
                'maior_mes_gasto': maior_mes_gasto,
            })

            if 'download' in request.GET:
                return gerar_relatorio_word_gastos(despesas, total_gastos, maior_gasto, maior_mes_gasto)

            return render(request, 'relatorio/relatorio3.html', contexto)
        else:
            return redirect('login')


# RELATORIO WORD GASTOS
def gerar_relatorio_word_gastos(despesas, total_gastos, maior_gasto, maior_mes_gasto):
    document = Document()
    document.add_heading('Relatório de Gastos', 0)

    document.add_heading('Lista de Despesas', level=1)
    for despesa in despesas:
        document.add_paragraph(f"{despesa['descricao']} - ${despesa['valor']} em {despesa['data_despesa']}")

    document.add_heading('Estatísticas de Gastos', level=1)
    document.add_paragraph(f"Total de Gastos: ${total_gastos}")
    if maior_gasto:
        document.add_paragraph(
            f"Maior Gasto: {maior_gasto['descricao']} - ${maior_gasto['valor']} em {maior_gasto['data_despesa']}")
    if maior_mes_gasto:
        document.add_paragraph(f"Maior Mês com Gasto: {maior_mes_gasto['data_despesa'].strftime('%B %Y')}")

    response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    response['Content-Disposition'] = 'attachment; filename=Relatorio_de_Gastos.docx'
    document.save(response)
    return response